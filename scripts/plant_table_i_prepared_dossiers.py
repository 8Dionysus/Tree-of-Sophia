#!/usr/bin/env python3
"""Plant supported prepared-dossier packages into the ToS philosophy tree.

This script reads the operator-local DOCX corpus and writes only structured
ToS surfaces: atlas indexes, branch bodies, source-anchor backlogs, and
pre-canon graph workbench rows.  The historical filename remains as a
compatibility entrypoint for the first Table I package.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from philosophy_multilingual_common import english_label, russian_label

REPO_ROOT = Path(__file__).resolve().parents[1]
DOC_ROOT = Path("/home/dionysus/Загрузки/ToS")
ATLAS_MANIFEST = REPO_ROOT / "ToS/philosophy/atlas/atlas.manifest.json"
PHILOSOPHY_MANIFEST = REPO_ROOT / "ToS/philosophy/philosophy.manifest.json"
DOSSIER_INDEX = REPO_ROOT / "ToS/philosophy/atlas/dossiers/index.jsonl"
DOSSIER_SUMMARY = REPO_ROOT / "ToS/philosophy/atlas/dossiers/graph-shape-summary.json"
DOSSIER_BRANCH_MANIFEST = REPO_ROOT / "ToS/philosophy/atlas/dossiers/branch.manifest.json"
PREPARED_DOSSIER_ROUTES = REPO_ROOT / "ToS/philosophy/atlas/dossiers/prepared-dossier-routes.json"
SOURCE_ANCHOR_BACKLOG = REPO_ROOT / "ToS/philosophy/atlas/dossiers/source-anchor-backlog.jsonl"
TERM_INDEX = REPO_ROOT / "ToS/philosophy/atlas/dossiers/term-index.jsonl"
TRANSMISSION_BACKLOG = REPO_ROOT / "ToS/philosophy/atlas/dossiers/transmission-backlog.jsonl"
LANGUAGE_PACKETS_MANIFEST = REPO_ROOT / "ToS/philosophy/graph-workbench/language-packets/branch.manifest.json"
LANGUAGE_PACKETS_README = REPO_ROOT / "ToS/philosophy/graph-workbench/language-packets/README.md"
LANGUAGE_PACKET_CONTRACT_REF = "ToS/philosophy/atlas/multilingual/text-bearing-nodes.contract.json"
LANGUAGE_REGISTRY_REF = "ToS/philosophy/atlas/multilingual/language-registry.json"
OBSOLETE_GENERATED_BRANCH = REPO_ROOT / "ToS/philosophy/eras/bronze-age/regions/ancient-near-east"
TEXT_BEARING_NODE_KINDS = {"text_corpus"}
INTAKE_RECORDED_ON = {"table-i": "2026-08-27", "table-ii": "2026-09-01", "table-iii": "2026-09-01"}


def load_packages() -> dict[str, dict[str, Any]]:
    payload = json.loads(PREPARED_DOSSIER_ROUTES.read_text(encoding="utf-8"))
    packages = payload.get("packages")
    if not isinstance(packages, dict):
        raise ValueError("prepared-dossier-routes.json must expose packages")
    if not all(isinstance(package, dict) for package in packages.values()):
        raise ValueError("prepared-dossier-routes.json packages must be objects")
    return packages


PACKAGES = load_packages()


def load_package_routes(table_id: str) -> dict[str, dict[str, Any]]:
    package = PACKAGES.get(table_id)
    if not isinstance(package, dict):
        raise ValueError(f"prepared-dossier-routes.json must expose a {table_id} package")
    routes = package.get("routes")
    if not isinstance(routes, list):
        raise ValueError(f"prepared-dossier-routes.json {table_id}.routes must be a list")
    defaults = package.get("route_defaults")
    if not routes:
        return {}
    if not isinstance(defaults, dict):
        raise ValueError(f"prepared-dossier-routes.json {table_id}.route_defaults must be an object")
    prepared_routes: dict[str, dict[str, Any]] = {}
    for route in routes:
        if not isinstance(route, dict):
            raise ValueError("prepared dossier routes must be objects")
        merged = {**defaults, **route}
        dossier_id = str(merged.get("dossier_id") or "")
        branch_path = str(merged.get("branch_path") or "")
        branch_role = str(merged.get("branch_role") or "")
        if not dossier_id or not branch_path or not branch_role:
            raise ValueError("prepared dossier routes must carry dossier_id, branch_path, and branch_role")
        if dossier_id in prepared_routes:
            raise ValueError(f"duplicate prepared dossier route: {dossier_id}")
        review_posture = str(merged.get("review_posture") or "")
        route_kind = str(merged.get("route_kind") or "")
        if not review_posture or not route_kind:
            raise ValueError("prepared dossier routes must resolve review_posture and route_kind")
        merged["table_id"] = table_id
        prepared_routes[dossier_id] = merged
    return prepared_routes


PACKAGE_ROUTES = {table_id: load_package_routes(table_id) for table_id in PACKAGES}
SUPPORTED_TABLES = tuple(table_id for table_id, routes in PACKAGE_ROUTES.items() if routes)
if "table-i" not in SUPPORTED_TABLES:
    raise ValueError("Table I compatibility package must remain supported")
ROUTES = {
    dossier_id: route
    for routes in PACKAGE_ROUTES.values()
    for dossier_id, route in routes.items()
}
if len(ROUTES) != sum(len(routes) for routes in PACKAGE_ROUTES.values()):
    raise ValueError("prepared dossier ids must be unique across supported packages")
TABLE_I_PACKAGE = PACKAGES["table-i"]
TABLE_II_PACKAGE = PACKAGES["table-ii"]
TABLE_I_ROUTES = PACKAGE_ROUTES["table-i"]
TABLE_II_ROUTES = PACKAGE_ROUTES["table-ii"]
BRANCHES = {
    dossier_id: (str(route["branch_path"]), str(route["branch_role"]))
    for dossier_id, route in TABLE_I_ROUTES.items()
}
TABLE_II_BRANCHES = {
    dossier_id: (str(route["branch_path"]), str(route["branch_role"]))
    for dossier_id, route in TABLE_II_ROUTES.items()
}
ALL_BRANCHES = {
    dossier_id: (str(route["branch_path"]), str(route["branch_role"]))
    for dossier_id, route in ROUTES.items()
}
TABLE_I_DOCX_SECTIONS = tuple(str(value) for value in TABLE_I_PACKAGE.get("docx_sections", []))
TABLE_II_DOCX_SECTIONS = tuple(str(value) for value in TABLE_II_PACKAGE.get("docx_sections", []))
if not TABLE_I_DOCX_SECTIONS:
    raise ValueError("prepared-dossier-routes.json table-i.docx_sections must be non-empty")


def master_rows_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/atlas/master-tables/{table_id}/rows.jsonl"


def master_manifest_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/atlas/master-tables/{table_id}/table.manifest.json"


def package_path(table_id: str, field: str) -> Path:
    value = PACKAGES[table_id].get(field)
    if not isinstance(value, str) or not value:
        raise ValueError(f"{table_id}.{field} must be a repo-relative path")
    return REPO_ROOT / value


def proposed_nodes_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/graph-workbench/proposed-nodes/{table_id}-prepared-dossiers.jsonl"


def proposed_relations_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/graph-workbench/proposed-relations/{table_id}-prepared-dossiers.jsonl"


def language_packets_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/graph-workbench/language-packets/{table_id}-text-bearing-nodes.jsonl"


def branch_fragments_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/graph-workbench/branch-fragments/{table_id}-prepared-dossier-branches.json"


def promotion_ledger_path(table_id: str) -> Path:
    return REPO_ROOT / f"ToS/philosophy/graph-workbench/promotion-ledger/{table_id}-prepared-dossiers.md"


# Historical names stay import-compatible for existing Table I tests and tools.
TABLE_I_ROWS = master_rows_path("table-i")
TABLE_I_MANIFEST = master_manifest_path("table-i")
INTAKE_MANIFEST = package_path("table-i", "intake_manifest_ref")
EXTRACTION_COVERAGE = package_path("table-i", "extraction_coverage_ref")
PROPOSED_NODES = proposed_nodes_path("table-i")
PROPOSED_RELATIONS = proposed_relations_path("table-i")
LANGUAGE_PACKETS = language_packets_path("table-i")
BRANCH_FRAGMENTS = branch_fragments_path("table-i")
PROMOTION_LEDGER = promotion_ledger_path("table-i")

NODE_TABLE = ("Node ID", "Тип узла", "Название", "Период", "Связи", "Приоритет")
RELATION_TABLE = ("Source node", "Relation", "Target node", "Комментарий", "Уверенность")
CORPUS_SOURCE_TABLE = ("Источник / корпус", "Тип", "Что даёт", "Доступ / где искать", "Надёжность")
CONTROL_SOURCE_TABLE = ("Источник", "Тип", "Зачем нужен", "Ограничения")
RISK_TABLES = {
    ("Проблема", "В чём риск", "Как контролировать в ToS", "Какие источники нужны"),
    ("Проблема", "Риск", "Как контролировать в ToS", "Какие источники нужны"),
    ("Риск", "Почему существенен", "Контроль ToS"),
}
TERM_TABLE = ("Термин", "Язык", "Транслитерация", "Краткое значение", "Роль в ToS")
INCOMING_TRANSMISSION_TABLE = (
    "Источник / предыдущий узел",
    "Что передано",
    "Канал передачи",
    "Уверенность",
    "Примечание",
)
OUTGOING_TRANSMISSION_TABLE = (
    "Следующий узел / эпоха",
    "Что передаётся",
    "Канал",
    "Уверенность",
    "Что проверить дальше",
)
IDENTITY_METADATA_TABLES = {("Поле", "Значение"), ("Поле", "Идентификация")}
METADATA_TABLES = IDENTITY_METADATA_TABLES | {("Параметр", "Значение")}
EXTRACTED_TABLE_FAMILIES = {
    NODE_TABLE: "proposed_nodes",
    RELATION_TABLE: "proposed_relations",
    CORPUS_SOURCE_TABLE: "corpus_or_edition_anchors",
    CONTROL_SOURCE_TABLE: "control_or_review_anchors",
    **{header: "risk_control_source_needs" for header in RISK_TABLES},
    TERM_TABLE: "terms",
    INCOMING_TRANSMISSION_TABLE: "incoming_transmissions",
    OUTGOING_TRANSMISSION_TABLE: "outgoing_transmissions",
}
DEFERRED_CONTEXT_FAMILIES = {
    ("Измерение", "Граница"): "boundary_dimensions",
    ("Язык / письменность / медиум", "Роль в строке", "Период", "Что сохранилось", "Риск"): "language_script_medium",
    ("Корпус / текст / артефакт", "Дата / слой", "Язык", "Жанр", "Сохранность", "Почему важен для ToS"): "corpora_texts_artifacts",
    ("Фигура / тип авторства", "Период", "Роль", "Связанные тексты", "Уверенность"): "figures_authorship",
    ("Фигура / тип авторства", "Период", "Роль", "Связанные тексты / объекты", "Уверенность"): "figures_authorship",
    ("Фигура / тип авторства", "Период", "Роль", "Связанные тексты / вещи", "Уверенность"): "figures_authorship",
    ("Жанр", "Функция", "Примеры", "Философская значимость"): "genres",
    ("Уровень", "Оценка"): "audit_levels",
}


def normalized_header_cell(value: str) -> str:
    value = scrub(value).lower().replace("ё", "е")
    value = re.sub(r"\s*/\s*", " / ", value)
    return re.sub(r"\s+", " ", value).strip()


def table_family(header: tuple[str, ...]) -> str:
    normalized = tuple(normalized_header_cell(value) for value in header)
    if (
        len(normalized) == 6
        and normalized[0] == "node id"
        and normalized[1] in {"тип узла", "тип"}
        and normalized[2:4] == ("название", "период")
        and normalized[4]
        in {
            "связи",
            "связи / функция",
            "основные связи",
            "ключевые связи",
        }
        and normalized[5] in {"приоритет", "приор."}
    ):
        return "proposed_nodes"
    if (
        len(normalized) == 5
        and normalized[:3]
        in {
            ("source node", "relation", "target node"),
            ("исходный узел", "отношение", "целевой узел"),
        }
        and normalized[3] == "комментарий"
        and normalized[4] in {"уверенность", "увер.", "ув."}
    ) or (
        len(normalized) == 6
        and normalized[:5] == ("edge id", "source", "relation", "target", "комментарий")
        and normalized[5] in {"уверенность", "увер."}
    ):
        return "proposed_relations"
    source_core = (
        normalized[1:]
        if normalized and normalized[0] in {"id", "код", "маркер"}
        else normalized
    )
    source_labels = {
        "источник / корпус",
        "источник",
        "корпус / архив",
        "корпус / портал",
        "корпус",
        "id / источник",
    }
    if (
        len(source_core) in {4, 5, 6}
        and source_core[0] in source_labels
        and (
            any(value.startswith("что дает") for value in source_core[1:])
            or (
                "дата / слой" in source_core
                and any(value.startswith("тип") for value in source_core[1:])
            )
        )
    ):
        return "corpus_or_edition_anchors"
    if (
        len(source_core) in {4, 5}
        and source_core[0] in {"источник", "id / источник"}
        and any(value.startswith("тип") for value in source_core[1:])
        and "зачем нужен" in source_core
        and any("ограничен" in value for value in source_core[1:])
    ):
        return "control_or_review_anchors"
    if (
        len(normalized) in {2, 3, 4}
        and normalized[0] in {"риск", "главный риск"}
        and any("контрол" in value for value in normalized[1:])
    ):
        return "risk_control_source_needs"
    if len(normalized) in {2, 3, 4} and normalized[0] in {"проблема", "риск"}:
        if len(normalized) == 2 and normalized[1] == "что контролировать":
            return "risk_control_source_needs"
        if len(normalized) == 3 and normalized[1] in {"почему существенен", "в чем ловушка"} and normalized[2] in {
            "контроль",
            "контроль tos",
            "контроль в tos",
        }:
            return "risk_control_source_needs"
        if len(normalized) == 4 and normalized[1] in {"в чем риск", "риск"} and normalized[2] in {
            "как контролировать в tos",
            "контроль в tos",
            "контроль tos",
        } and normalized[3] in {"какие источники нужны", "нужные источники", "что требуется"}:
            return "risk_control_source_needs"
    if len(normalized) == 5 and normalized[:2] == ("термин", "язык") and normalized[2].startswith("транслитерация") and normalized[3] in {
        "краткое значение",
        "значение",
    } and normalized[4] == "роль в tos":
        return "terms"
    if len(normalized) == 5 and normalized[0] in {
        "источник / предыдущий узел",
        "источник / previous node",
    } and normalized[1] == "что передано" and normalized[2] in {
        "канал передачи",
        "канал",
    } and normalized[3:] == ("уверенность", "примечание"):
        return "incoming_transmissions"
    if (
        len(normalized) == 5
        and normalized[:4]
        == ("следующий узел / эпоха", "что передается", "канал", "уверенность")
        and normalized[4] in {"что проверить дальше", "проверить дальше"}
    ):
        return "outgoing_transmissions"
    if len(normalized) == 2 and normalized[0] in {"поле", "параметр"} and normalized[1] in {
        "значение",
        "идентификация",
    }:
        return "dossier_identity_metadata" if normalized[0] == "поле" else "dossier_identity_metadata_alias"
    if (
        len(normalized) == 6
        and normalized[0] in {
            "корпус / текст",
            "корпус / текст / артефакт",
            "корпус / артефакт",
        }
        and normalized[1] in {"дата / слой", "дата"}
        and normalized[2] == "язык"
        and normalized[3] in {"жанр", "жанр / жанры"}
        and normalized[4] == "сохранность"
        and normalized[5] in {"почему важен для tos", "значение tos", "tos-функция"}
    ):
        return "corpora_texts_artifacts"
    if (
        len(normalized) == 5
        and normalized[0] in {"фигура / тип авторства", "фигура / тип"}
        and normalized[1:3] == ("период", "роль")
        and (
            normalized[3] == "связанные тексты"
            or normalized[3].startswith("связанные тексты / ")
        )
        and normalized[4] == "уверенность"
    ):
        return "figures_authorship"
    if (
        len(normalized) == 5
        and normalized[0] in {"язык / письменность / медиум", "язык / письмо / медиум"}
        and normalized[1] in {"роль в строке", "роль"}
        and normalized[2:] == ("период", "что сохранилось", "риск")
    ):
        return "language_script_medium"
    if len(normalized) == 2 and normalized[0] in {"уровень", "уровень оценки"} and normalized[1] in {
        "оценка",
        "вывод",
    }:
        return "audit_levels"
    for deferred_header, family in DEFERRED_CONTEXT_FAMILIES.items():
        if normalized == tuple(normalized_header_cell(value) for value in deferred_header):
            return family
    return "other_context"


STRUCTURED_FAMILIES = {
    "proposed_nodes",
    "proposed_relations",
    "corpus_or_edition_anchors",
    "control_or_review_anchors",
    "risk_control_source_needs",
    "terms",
    "incoming_transmissions",
    "outgoing_transmissions",
}


def row_value(row: dict[str, str], *aliases: str) -> str:
    normalized = {normalized_header_cell(key): value for key, value in row.items()}
    for alias in aliases:
        value = normalized.get(normalized_header_cell(alias))
        if value:
            return value
    return ""


@dataclass
class Dossier:
    table_id: str
    dossier_id: str
    title: str
    source_document: str
    docx_path: Path
    docx_section: str
    paragraph_count: int
    table_count: int
    table_row: str
    master_table: str
    master_status: str
    master_confidence: str
    branch_path: str | None
    branch_role: str | None
    admission_status: str
    identity_diagnostics: list[str]
    node_rows: list[dict[str, Any]]
    relation_rows: list[dict[str, Any]]
    source_rows: list[dict[str, Any]]
    term_rows: list[dict[str, Any]]
    transmission_rows: list[dict[str, Any]]
    metadata_identity_posture: str
    metadata_headers: list[list[str]]
    coverage_tables: list[dict[str, Any]]


def repo_ref(path: Path) -> str:
    return path.relative_to(REPO_ROOT).as_posix()


def scrub(text: str) -> str:
    text = re.sub(r"\ue200filecite\ue202[^ \n\t]*\ue201", "", text)
    text = re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()
    return text


def slugify(value: str) -> str:
    value = value.lower()
    value = value.replace("ё", "е")
    value = re.sub(r"[^a-z0-9а-я]+", "-", value)
    return value.strip("-") or "unnamed"


def branch_id_for(path_ref: str) -> str:
    suffix = path_ref.removeprefix("ToS/philosophy/").replace("/", ".")
    return f"philosophy.{suffix}"


def load_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"{repo_ref(path)} must contain a JSON object")
    return payload


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def load_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line))
    return rows


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    rendered = "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in rows)
    path.write_text(rendered, encoding="utf-8")


def row_dict(headers: tuple[str, ...], cells: list[str]) -> dict[str, str]:
    return {headers[index]: scrub(cells[index]) if index < len(cells) else "" for index in range(len(headers))}


def table_header(table: Any) -> tuple[str, ...]:
    if not table.rows:
        return ()
    return tuple(scrub(cell.text) for cell in table.rows[0].cells)


def table_body(table: Any) -> list[list[str]]:
    body: list[list[str]] = []
    for row in table.rows[1:]:
        cells = [scrub(cell.text) for cell in row.cells]
        if any(cells):
            body.append(cells)
    return body


DOSSIER_ID_PATTERN = re.compile(r"(?<![A-Za-z0-9])(?:A\d{2}|T[23]-\d{2})(?!\d)")


def normalize_row_to_expand(value: str, dossier_id: str) -> str | None:
    """Resolve a ROW_TO_EXPAND value without accepting a different full id."""
    match = DOSSIER_ID_PATTERN.search(scrub(value))
    if match:
        return match.group(0)
    number = re.fullmatch(r"\s*(\d{1,2})(?:\s*\([^)]*\))?\s*", scrub(value))
    if number:
        prefix = dossier_id.split("-", 1)[0]
        return f"{prefix}-{int(number.group(1)):02d}"
    return None


def extract_dossier_id(path: Path) -> str:
    match = DOSSIER_ID_PATTERN.search(path.name)
    if not match:
        raise ValueError(f"Cannot find dossier id in {path}")
    return match.group(0)


def blocked_dossiers(table_id: str) -> dict[str, dict[str, Any]]:
    values = PACKAGES[table_id].get("blocked_dossiers", [])
    if not isinstance(values, list):
        raise ValueError(f"{table_id}.blocked_dossiers must be a list")
    result: dict[str, dict[str, Any]] = {}
    for value in values:
        if not isinstance(value, dict) or not value.get("dossier_id"):
            raise ValueError(f"{table_id}.blocked_dossiers entries must carry dossier_id")
        dossier_id = str(value["dossier_id"])
        if dossier_id in result or dossier_id in PACKAGE_ROUTES[table_id]:
            raise ValueError(f"duplicate or routed blocked dossier id: {dossier_id}")
        result[dossier_id] = value
    return result


def discover_docx(table_id: str = "table-i") -> list[Path]:
    sections = tuple(str(value) for value in PACKAGES[table_id].get("docx_sections", []))
    paths = sorted(
        path
        for section in sections
        for path in (DOC_ROOT / str(PACKAGES[table_id].get("docx_root") or "") / section).glob("*.docx")
        if DOSSIER_ID_PATTERN.search(path.name)
    )
    ids = [extract_dossier_id(path) for path in paths]
    expected = sorted(set(PACKAGE_ROUTES[table_id]) | set(blocked_dossiers(table_id)))
    if len(ids) != len(set(ids)):
        raise SystemExit(f"Prepared dossier ids must be unique, found {sorted(ids)}")
    if sorted(ids) != expected:
        raise SystemExit(f"Expected {table_id} prepared dossier ids {expected}, found {sorted(ids)}")
    return paths


def load_docx_document(path: Path) -> Any:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("Install python-docx to plant prepared DOCX dossiers.") from exc
    return Document(path)


def clean_dossier_title(title: str, dossier_id: str) -> str:
    value = re.sub(r"^ToS Deep Research\s*[:_—-]*\s*", "", scrub(title), flags=re.IGNORECASE)
    value = re.sub(rf"^{re.escape(dossier_id)}\s*[:—-]*\s*", "", value, flags=re.IGNORECASE)
    return value.strip()


def normalized_title(value: str) -> str:
    value = value.lower().replace("ё", "е").replace("ā", "a")
    value = re.sub(r"[^0-9a-zа-я]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def parse_dossier(path: Path, master_row: dict[str, Any], table_id: str = "table-i") -> Dossier:
    dossier_id = extract_dossier_id(path)
    if master_row.get("row_id") != dossier_id or master_row.get("table_id") != table_id:
        raise ValueError(f"{dossier_id} does not match its {table_id} master row")
    route = PACKAGE_ROUTES[table_id].get(dossier_id)
    blocked = blocked_dossiers(table_id).get(dossier_id)
    if (route is None) == (blocked is None):
        raise ValueError(f"{dossier_id} must be either routed or explicitly blocked")
    branch_path = str(route["branch_path"]) if route else None
    branch_role = str(route["branch_role"]) if route else None
    admission_status = "admitted" if route else str(blocked.get("posture") or "blocked")
    document = load_docx_document(path)
    paragraphs = [scrub(paragraph.text) for paragraph in document.paragraphs if scrub(paragraph.text)]
    title = paragraphs[0] if paragraphs else f"ToS Deep Research: {dossier_id}"
    node_rows: list[dict[str, Any]] = []
    relation_rows: list[dict[str, Any]] = []
    source_rows: list[dict[str, Any]] = []
    term_rows: list[dict[str, Any]] = []
    transmission_rows: list[dict[str, Any]] = []
    metadata_headers: list[list[str]] = []
    coverage_tables: list[dict[str, Any]] = []
    node_label_overrides = route.get("reviewed_node_label_overrides", {}) if route else {}
    if not isinstance(node_label_overrides, dict):
        raise ValueError(f"{dossier_id} reviewed_node_label_overrides must be an object")
    applied_node_label_overrides: set[str] = set()
    observed_table_value: str | None = None
    observed_row_value: str | None = None
    table_row = dossier_id
    master_table = {"table-i": "I", "table-ii": "II", "table-iii": "III"}.get(table_id, table_id)
    identity_diagnostics: list[str] = []

    for table_index, table in enumerate(document.tables, 1):
        header = table_header(table)
        rows = table_body(table)
        family = table_family(header)
        if blocked:
            coverage_class = "quarantined_identity_mismatch"
            coverage_family = "blocked_master_identity_mismatch"
        elif family in STRUCTURED_FAMILIES:
            coverage_class = "structured_primary_extracted"
            coverage_family = family
        elif family == "dossier_identity_metadata":
            coverage_class = "identity_metadata_examined"
            coverage_family = "dossier_identity_metadata"
        else:
            coverage_class = "deferred_context"
            coverage_family = "other_context" if family == "dossier_identity_metadata_alias" else family
        coverage_tables.append(
            {
                "coverage_class": coverage_class,
                "family": coverage_family,
                "header": list(header),
                **({"underlying_family": family} if blocked else {}),
                "row_count": len(rows),
                "source_table_index": table_index,
            }
        )
        if family in {"dossier_identity_metadata", "dossier_identity_metadata_alias"}:
            metadata_headers.append(list(header))
            for cells in rows:
                row = row_dict(header, cells)
                field_name = row_value(row, "Поле", "Параметр")
                field_value = row_value(row, "Значение", "Идентификация")
                if field_name == "ROW_TO_EXPAND":
                    observed_row_value = field_value or observed_row_value
                if field_name == "Таблица":
                    observed_table_value = field_value or observed_table_value
        if blocked:
            continue
        if family == "proposed_nodes":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                original_id = row_value(row, "Node ID") or f"{dossier_id}-node-{row_index:03d}"
                candidate_id = f"{table_id}-{dossier_id.lower()}-node-{row_index:03d}"
                node_kind_label = row_value(row, "Тип узла", "Тип")
                node_kind = (
                    slugify(node_kind_label).replace("-", "_")
                    if node_kind_label
                    else "unspecified"
                )
                label = row_value(row, "Название") or original_id
                label_override = node_label_overrides.get(original_id)
                override_projection: dict[str, str] = {}
                if label_override is not None:
                    if not isinstance(label_override, dict):
                        raise ValueError(f"{dossier_id} node label override for {original_id} must be an object")
                    reviewed_label = scrub(str(label_override.get("label") or ""))
                    review_reason = scrub(str(label_override.get("reason") or ""))
                    if not reviewed_label or not review_reason:
                        raise ValueError(
                            f"{dossier_id} node label override for {original_id} requires label and reason"
                        )
                    label = reviewed_label
                    applied_node_label_overrides.add(original_id)
                    override_projection = {
                        "label_normalization": "reviewed_source_punctuation_normalization",
                        "label_override_reason": review_reason,
                    }
                node_rows.append(
                    {
                        "atlas_row_id": dossier_id,
                        "authority_posture": "prepared_research_candidate",
                        "branch_path": branch_path,
                        "candidate_id": candidate_id,
                        "canon_status": "pre-canon",
                        "connections": row_value(
                            row,
                            "Связи",
                            "Связи / функция",
                            "Основные связи",
                            "Ключевые связи",
                        ),
                        "dossier_id": dossier_id,
                        "label": label,
                        **override_projection,
                        "node_kind": node_kind,
                        "original_node_id": original_id,
                        "period": row_value(row, "Период"),
                        "priority": row_value(row, "Приоритет", "Приор."),
                        "source_document": path.name,
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        "source_ref": repo_ref(proposed_nodes_path(table_id)),
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                    }
                )
        elif family == "proposed_relations":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                relation_label = row_value(row, "Relation", "Отношение")
                relation_kind = slugify(relation_label).replace("-", "_")
                candidate_id = f"{table_id}-{dossier_id.lower()}-relation-{len(relation_rows) + 1:03d}"
                relation_rows.append(
                    {
                        "atlas_row_id": dossier_id,
                        "authority_posture": "prepared_research_candidate",
                        "branch_path": branch_path,
                        "candidate_id": candidate_id,
                        "canon_status": "pre-canon",
                        "comment": row_value(row, "Комментарий"),
                        "confidence": row_value(row, "Уверенность", "Увер.", "Ув."),
                        "dossier_id": dossier_id,
                        "relation_kind": relation_kind or "related_to",
                        "relation_label": relation_label,
                        "source_document": path.name,
                        "source_endpoint_label": row_value(
                            row,
                            "Source node",
                            "Source",
                            "Исходный узел",
                        ),
                        "source_ref": repo_ref(proposed_relations_path(table_id)),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                        "target_endpoint_label": row_value(
                            row,
                            "Target node",
                            "Target",
                            "Целевой узел",
                        ),
                    }
                )
        elif family == "corpus_or_edition_anchors":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                source_rows.append(
                    {
                        "anchor_kind": "corpus_or_edition_anchor",
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "contribution": row_value(
                            row,
                            "Что даёт",
                            "Что дает",
                            "Что даёт ToS",
                            "Что дает ToS",
                            "Что даёт / где искать",
                            "Что дает / где искать",
                            "Что даёт / доступ",
                            "Что дает / доступ",
                        ),
                        "dossier_id": dossier_id,
                        "reliability": row_value(
                            row,
                            "Надёжность",
                            "Надежность",
                            "Надёжность / ограничение",
                            "Надежность / ограничение",
                            "Надёжность / ограничения",
                            "Надежность / ограничения",
                            "Надёжность и ограничения",
                            "Надежность и ограничения",
                            "Надёжность / caveat",
                            "Надежность / caveat",
                            "Ограничения",
                            "Доступ / надёжность",
                            "Доступ / надежность",
                        ),
                        "route_status": "source_anchor_backlog",
                        "source_access": row_value(
                            row,
                            "Доступ / где искать",
                            "Доступ",
                            "Ссылка",
                            "URL / DOI",
                            "Доступ / stable URL",
                            "Доступ / надёжность",
                            "Доступ / надежность",
                            "Доступ / замечание",
                        ),
                        "source_document": path.name,
                        "source_label": row_value(
                            row,
                            "Источник / корпус",
                            "Источник",
                            "Корпус / архив",
                            "Корпус / портал",
                            "Корпус",
                            "ID / источник",
                        ),
                        "source_ref": repo_ref(SOURCE_ANCHOR_BACKLOG),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        "source_type": row_value(row, "Тип", "Тип / дата", "Тип / содержание"),
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                    }
                )
        elif family == "control_or_review_anchors":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                source_access = row_value(
                    row,
                    "Доступ / где искать",
                    "Доступ",
                    "Ссылка",
                    "URL / DOI",
                    "Доступ / stable URL",
                    "Ограничения / доступ",
                )
                source_rows.append(
                    {
                        "anchor_kind": "control_or_review_anchor",
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "contribution": row_value(row, "Зачем нужен"),
                        "dossier_id": dossier_id,
                        "limitations": row_value(
                            row,
                            "Ограничения",
                            "Ограничение",
                            "Ограничения / контроль",
                            "Ограничения / доступ",
                        ),
                        "route_status": "source_anchor_backlog",
                        "source_document": path.name,
                        "source_label": row_value(row, "Источник", "ID / источник"),
                        "source_ref": repo_ref(SOURCE_ANCHOR_BACKLOG),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        "source_type": row_value(row, "Тип", "Тип / дата", "Тип / содержание"),
                        **({"source_access": source_access} if source_access else {}),
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                    }
                )
        elif family == "risk_control_source_needs":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                first_header = normalized_header_cell(header[0]) if header else ""
                problem = row_value(row, "Проблема")
                risk = row_value(row, "В чём риск", "Почему существенен", "В чём ловушка")
                if first_header in {"риск", "главный риск"}:
                    problem = ""
                    risk = risk or row_value(row, "Риск", "Главный риск")
                else:
                    risk = risk or row_value(row, "Риск")
                source_rows.append(
                    {
                        "anchor_kind": "risk_control_source_need",
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "control": row_value(
                            row,
                            "Как контролировать в ToS",
                            "Контроль в ToS",
                            "Контроль ToS",
                            "Контроль",
                            "Что контролировать",
                            "Что именно контролировать",
                            "Контрольный принцип",
                        ),
                        "dossier_id": dossier_id,
                        "needed_sources": row_value(row, "Какие источники нужны", "Нужные источники", "Что требуется"),
                        "problem": problem,
                        "risk": risk,
                        "route_status": "source_anchor_backlog",
                        "source_document": path.name,
                        "source_ref": repo_ref(SOURCE_ANCHOR_BACKLOG),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                    }
                )
        elif family == "terms":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                term_rows.append(
                    {
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "dossier_id": dossier_id,
                        "language": row_value(row, "Язык"),
                        "meaning": row_value(row, "Краткое значение", "Значение"),
                        "role_in_tos": row_value(row, "Роль в ToS"),
                        "source_document": path.name,
                        "source_ref": repo_ref(TERM_INDEX),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                        "term": row_value(row, "Термин"),
                        "transliteration": row_value(
                            row,
                            "Транслитерация",
                            "Транслитерация / форма",
                            "Транслитерация / перевод",
                            "Транслитерация / аббр.",
                        ),
                    }
                )
        elif family == "incoming_transmissions":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                transmission_rows.append(
                    {
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "confidence": row_value(row, "Уверенность"),
                        "dossier_id": dossier_id,
                        "direction": "incoming",
                        "from_or_to": row_value(
                            row,
                            "Источник / предыдущий узел",
                            "Источник / previous node",
                        ),
                        "note": row_value(row, "Примечание"),
                        "source_document": path.name,
                        "source_ref": repo_ref(TRANSMISSION_BACKLOG),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                        "transmission_channel": row_value(row, "Канал передачи", "Канал"),
                        "transmitted": row_value(row, "Что передано"),
                    }
                )
        elif family == "outgoing_transmissions":
            for row_index, cells in enumerate(rows, 1):
                row = row_dict(header, cells)
                transmission_rows.append(
                    {
                        "atlas_row_id": dossier_id,
                        "branch_path": branch_path,
                        "confidence": row_value(row, "Уверенность"),
                        "dossier_id": dossier_id,
                        "direction": "outgoing",
                        "from_or_to": row_value(row, "Следующий узел / эпоха"),
                        "source_document": path.name,
                        "source_ref": repo_ref(TRANSMISSION_BACKLOG),
                        "source_row_index": row_index,
                        "source_table_index": table_index,
                        **({"table_id": table_id} if table_id != "table-i" else {}),
                        "transmission_channel": row_value(row, "Канал"),
                        "transmitted": row_value(row, "Что передаётся", "Что передается"),
                        "verify_next": row_value(row, "Что проверить дальше", "Проверить дальше"),
                    }
                )

    unapplied_node_label_overrides = set(node_label_overrides) - applied_node_label_overrides
    if unapplied_node_label_overrides:
        raise ValueError(
            f"{dossier_id} node label overrides did not match extracted node ids: "
            f"{sorted(unapplied_node_label_overrides)}"
        )

    if observed_row_value:
        observed_id = normalize_row_to_expand(observed_row_value, dossier_id)
        if observed_id != dossier_id:
            raise ValueError(f"{dossier_id} DOCX ROW_TO_EXPAND does not match its master-table identity")
        table_row = observed_row_value
    if observed_table_value:
        normalized_table = scrub(observed_table_value)
        accepted_tables = {
            "table-i": {"I", "Table I", "Таблица I"},
            "table-ii": {"II", "Table II", "Таблица II"},
            "table-iii": {
                "III",
                "Table III",
                "Таблица III",
                "III — модерность и современность",
                "III — модерность и современность как сеть письменных инфраструктур",
                "Таблица III — модерность и современность",
            },
        }.get(table_id, {table_id})
        if normalized_table not in accepted_tables:
            raise ValueError(f"{dossier_id} DOCX table identity is not {table_id}: {normalized_table}")
    if observed_row_value and observed_table_value:
        metadata_identity_posture = "docx_metadata_cross_checked"
    elif observed_row_value or observed_table_value:
        metadata_identity_posture = "partial_docx_metadata_cross_checked"
    else:
        metadata_identity_posture = "master_table_identity_fallback"

    normalized_master = master_row.get("normalized")
    if not isinstance(normalized_master, dict):
        raise ValueError(f"{dossier_id} master row must expose normalized metadata")
    if table_id == "table-i" and not observed_row_value:
        observed_title_id = DOSSIER_ID_PATTERN.search(title)
        if observed_title_id and observed_title_id.group(0) != dossier_id:
            raise ValueError(
                f"{dossier_id} DOCX title does not match its master-table identity: "
                f"{observed_title_id.group(0)}"
            )
        if observed_title_id is None:
            expected_title = str(route.get("accepted_input_title") or "")
            if not expected_title or normalized_title(title) != normalized_title(expected_title):
                raise ValueError(
                    f"{dossier_id} DOCX title does not match its reviewed Table I route: "
                    f"{title!r} != {expected_title!r}"
                )
            metadata_identity_posture = "filename_title_reviewed_route_cross_checked"
    if table_id in {"table-ii", "table-iii"}:
        observed_clean_title = clean_dossier_title(title, dossier_id)
        if blocked:
            expected_input_title = str(blocked.get("observed_input_title") or "")
            if normalized_title(observed_clean_title) != normalized_title(expected_input_title):
                raise ValueError(f"{dossier_id} blocked artifact title changed; review quarantine before planting")
            identity_diagnostics.append("master_identity_mismatch_quarantined")
            metadata_identity_posture = "filename_and_artifact_title_recorded_master_mismatch"
        else:
            if "accepted_input_title" in route:
                expected_title = scrub(str(route["accepted_input_title"]))
                if not normalized_title(expected_title):
                    raise ValueError(f"{dossier_id} accepted_input_title must be non-empty")
            else:
                expected_title = str(normalized_master.get("research_node") or "")
            normalized_expected_title = normalized_title(expected_title)
            body_text = normalized_title(" ".join(paragraphs[1:]))
            title_matches = normalized_title(observed_clean_title) == normalized_expected_title
            title_matches = title_matches or normalized_expected_title in body_text
            if not title_matches:
                raise ValueError(
                    f"{dossier_id} DOCX title does not match its reviewed {table_id} route: "
                    f"{observed_clean_title!r} != {expected_title!r}"
                )
            if not observed_row_value and not observed_table_value:
                metadata_identity_posture = "filename_title_master_cross_checked"

    if route:
        route_projection = explicit_route_fields(
            dossier_id,
            master_status=str(normalized_master.get("status") or ""),
            master_confidence=str(normalized_master.get("confidence") or ""),
        )
        for extracted_rows in (node_rows, relation_rows, source_rows, term_rows, transmission_rows):
            for row in extracted_rows:
                row.update(route_projection)

    return Dossier(
        table_id=table_id,
        dossier_id=dossier_id,
        title=title,
        source_document=path.name,
        docx_path=path,
        docx_section=path.parent.name,
        paragraph_count=len(paragraphs),
        table_count=len(document.tables),
        table_row=table_row,
        master_table=master_table,
        master_status=str(normalized_master.get("status") or ""),
        master_confidence=str(normalized_master.get("confidence") or ""),
        branch_path=branch_path,
        branch_role=branch_role,
        admission_status=admission_status,
        identity_diagnostics=identity_diagnostics,
        node_rows=node_rows,
        relation_rows=relation_rows,
        source_rows=source_rows,
        term_rows=term_rows,
        transmission_rows=transmission_rows,
        metadata_identity_posture=metadata_identity_posture,
        metadata_headers=metadata_headers,
        coverage_tables=coverage_tables,
    )


def route_metadata(dossier_id: str) -> dict[str, Any]:
    route = ROUTES.get(dossier_id)
    if not isinstance(route, dict):
        raise ValueError(f"missing prepared dossier route for {dossier_id}")
    return route


def explicit_route_fields(
    dossier_id: str,
    *,
    master_status: str = "",
    master_confidence: str = "",
) -> dict[str, Any]:
    """Return only route metadata that changes the default planting contract."""
    route = route_metadata(dossier_id)
    table_id = str(route["table_id"])
    package = PACKAGES[table_id]
    defaults = package["route_defaults"]
    review_policy = package.get("review_policy")
    if not isinstance(review_policy, dict):
        review_policy = {}
    manual_statuses = {str(value) for value in review_policy.get("manual_review_statuses", [])}
    max_confidence = review_policy.get("manual_review_max_confidence")
    try:
        confidence_value = int(master_confidence)
    except (TypeError, ValueError):
        confidence_value = None
    manual_by_policy = master_status in manual_statuses or (
        isinstance(max_confidence, int)
        and confidence_value is not None
        and confidence_value <= max_confidence
    )
    manual_review = route["review_posture"] == "manual_review_required" or manual_by_policy
    non_default_route = route["route_kind"] != defaults["route_kind"]
    if not manual_review and not non_default_route:
        return {}

    payload: dict[str, Any] = {
        "review_posture": "manual_review_required" if manual_review else str(route["review_posture"]),
        "route_kind": str(route["route_kind"]),
    }
    if manual_review:
        payload.update(
            {
                "master_confidence": master_confidence,
                "master_status": master_status,
            }
        )
    if route.get("review_reason"):
        payload["review_reason"] = str(route["review_reason"])
    elif manual_by_policy:
        payload["review_reason"] = (
            f"{table_id} master status {master_status or 'unknown'} and confidence "
            f"{master_confidence or 'unknown'} require manual review under the package review policy"
        )
    if isinstance(route.get("route_constraints"), list):
        payload["route_constraints"] = [str(value) for value in route["route_constraints"]]
    return payload


def dossier_sort_key(dossier: Dossier) -> tuple[str, int]:
    match = re.search(r"(\d+)$", dossier.dossier_id)
    return dossier.table_id, int(match.group(1)) if match else 0


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _xml_text(root: ElementTree.Element, path: str, namespaces: dict[str, str]) -> str | None:
    element = root.find(path, namespaces)
    if element is None or element.text is None:
        return None
    value = scrub(element.text)
    return value or None


def docx_package_metadata(path: Path) -> dict[str, Any]:
    custom_properties: dict[str, str] = {}
    creator: str | None = None
    last_modified_by: str | None = None
    signature_parts: list[str] = []
    with zipfile.ZipFile(path) as package:
        names = package.namelist()
        signature_parts = sorted(name for name in names if name.startswith("_xmlsignatures/"))
        if "docProps/custom.xml" in names:
            custom_root = ElementTree.fromstring(package.read("docProps/custom.xml"))
            for property_element in custom_root:
                name = property_element.attrib.get("name")
                if not name:
                    continue
                value = scrub("".join(property_element.itertext()))
                if value:
                    custom_properties[name] = value
        if "docProps/core.xml" in names:
            core_root = ElementTree.fromstring(package.read("docProps/core.xml"))
            creator = _xml_text(core_root, "dc:creator", {"dc": "http://purl.org/dc/elements/1.1/"})
            last_modified_by = _xml_text(
                core_root,
                "cp:lastModifiedBy",
                {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"},
            )
    return {
        "creator": creator,
        "custom_generator": custom_properties.get("generator"),
        "last_modified_by": last_modified_by,
        "signature_part_count": len(signature_parts),
    }


def intake_fingerprint(records: list[dict[str, Any]]) -> str:
    body = "".join(
        f"{row['relative_path']}\t{row['size_bytes']}\t{row['sha256']}\n"
        for row in sorted(records, key=lambda item: str(item["relative_path"]))
    )
    return hashlib.sha256(body.encode("utf-8")).hexdigest()


def build_intake_manifest(dossiers: list[Dossier]) -> dict[str, Any]:
    table_ids = {dossier.table_id for dossier in dossiers}
    if len(table_ids) != 1:
        raise ValueError("intake manifest must be built for exactly one table package")
    table_id = next(iter(table_ids))
    intake_manifest = package_path(table_id, "intake_manifest_ref")
    records: list[dict[str, Any]] = []
    for dossier in sorted(dossiers, key=dossier_sort_key):
        package_metadata = docx_package_metadata(dossier.docx_path)
        records.append(
            {
                "creator": package_metadata["creator"],
                "custom_generator": package_metadata["custom_generator"],
                "dossier_id": dossier.dossier_id,
                **({"admission_status": dossier.admission_status} if table_id != "table-i" else {}),
                "last_modified_by": package_metadata["last_modified_by"],
                "relative_path": f"{dossier.docx_section}/{dossier.source_document}",
                "section": dossier.docx_section,
                "sha256": file_sha256(dossier.docx_path),
                "signature_part_count": package_metadata["signature_part_count"],
                "size_bytes": dossier.docx_path.stat().st_size,
            }
        )
    by_section = {
        section: [row for row in records if row["section"] == section]
        for section in PACKAGES[table_id].get("docx_sections", [])
    }
    return {
        "schema_version": "tos_philosophy_docx_intake_manifest_v1",
        "path": repo_ref(intake_manifest),
        "owner_repo": "Tree-of-Sophia",
        "owner_surface": "ToS/research-packets/deep-research/philosophy/packet-contract.md",
        "recorded_on": INTAKE_RECORDED_ON[table_id],
        **({"table_id": table_id} if table_id != "table-i" else {}),
        "artifact_role": (
            "operator-local prepared Table I dossier extraction input"
            if table_id == "table-i"
            else f"operator-local prepared {table_id} dossier extraction input"
        ),
        "capture_posture": {
            "custody": "operator_local_untracked_payload",
            "origin_verification": "unverified",
            "author_identity": None,
            "session_or_export_id": None,
            "generator_property_observed": sorted(
                {str(row["custom_generator"]) for row in records if row.get("custom_generator")}
            ),
            "signature_posture": (
                "signature_parts_present"
                if any(int(row["signature_part_count"]) for row in records)
                else "no_ooxml_signature_parts_observed"
            ),
        },
        **(
            {
                "artifact_trust_posture": {
                    "artifact_class": "operator_local_research_docx_bundle",
                    "registered_trust_class": False,
                    "verdict": "unknown",
                    "reason": "No registered research-DOCX trust class, authenticated producer, signature, or verified origin is available; exact local research intake remains claim-limited.",
                }
            }
            if table_id != "table-i"
            else {}
        ),
        "fingerprint_contract": {
            "algorithm": "sha256",
            "record_format": "relative_path<TAB>size_bytes<TAB>sha256<LF>",
            "ordering": "relative_path_utf8_ascending",
        },
        "bundle_fingerprint": intake_fingerprint(records),
        "section_fingerprints": {
            section: intake_fingerprint(section_records)
            for section, section_records in by_section.items()
        },
        "section_counts": {section: len(section_records) for section, section_records in by_section.items()},
        "file_count": len(records),
        **(
            {
                "admitted_file_count": sum(1 for dossier in dossiers if dossier.admission_status == "admitted"),
                "quarantined_file_count": sum(1 for dossier in dossiers if dossier.admission_status != "admitted"),
            }
            if table_id != "table-i"
            else {}
        ),
        "files": records,
        "claim_limit": (
            "This manifest proves only the bytes, sizes, logical section paths, and OOXML metadata observed in the "
            "operator-local capture at planting time. It does not prove authorship, export-session identity, origin, "
            "signature trust, source-witness status, claim truth, rights, review, doctrine, or canon acceptance."
        ),
    }


def coverage_summary(dossiers: list[Dossier]) -> dict[str, Any]:
    class_counts: Counter[str] = Counter()
    family_counts: Counter[str] = Counter()
    underlying_family_counts: Counter[str] = Counter()
    header_counts: Counter[tuple[str, ...]] = Counter()
    for dossier in dossiers:
        for table in dossier.coverage_tables:
            row_count = int(table["row_count"])
            class_counts[str(table["coverage_class"])] += row_count
            family_counts[str(table["family"])] += row_count
            if table.get("underlying_family"):
                underlying_family_counts[str(table["underlying_family"])] += row_count
            header_counts[tuple(str(value) for value in table["header"])] += row_count
    return {
        "dossier_count": len(dossiers),
        "table_body_row_count": sum(class_counts.values()),
        "coverage_class_counts": dict(sorted(class_counts.items())),
        "family_row_counts": dict(sorted(family_counts.items())),
        **(
            {"underlying_family_row_counts": dict(sorted(underlying_family_counts.items()))}
            if underlying_family_counts
            else {}
        ),
        "headers": [
            {"header": list(header), "row_count": count}
            for header, count in sorted(header_counts.items(), key=lambda item: (item[0], item[1]))
        ],
    }


def build_extraction_coverage(dossiers: list[Dossier]) -> dict[str, Any]:
    table_ids = {dossier.table_id for dossier in dossiers}
    if len(table_ids) != 1:
        raise ValueError("extraction coverage must be built for exactly one table package")
    table_id = next(iter(table_ids))
    intake_manifest = package_path(table_id, "intake_manifest_ref")
    extraction_coverage = package_path(table_id, "extraction_coverage_ref")
    diagnostics: list[dict[str, Any]] = []
    dossier_rows: list[dict[str, Any]] = []
    for dossier in sorted(dossiers, key=dossier_sort_key):
        summary = coverage_summary([dossier])
        risk_rows = int(summary["family_row_counts"].get("risk_control_source_needs", 0))
        dossier_diagnostics: list[str] = list(dossier.identity_diagnostics)
        if dossier.admission_status != "admitted":
            diagnostics.append(
                {
                    "code": dossier.admission_status,
                    "dossier_id": dossier.dossier_id,
                    "posture": "artifact_accounted_but_no_structured_semantic_output",
                    "message": "The supplied artifact conflicts with its master identity and is quarantined from planting.",
                }
            )
        elif dossier.table_count == 0:
            dossier_diagnostics.extend(["structured_tables_absent", "structured_risk_table_absent"])
            diagnostics.append(
                {
                    "code": "structured_tables_absent",
                    "dossier_id": dossier.dossier_id,
                    "posture": "prose_only_artifact_no_structured_semantic_projection",
                    "message": "The admitted DOCX contains no tables; its bytes and paragraph count are captured, but no nodes, relations, source anchors, terms, or transmissions are synthesized from prose.",
                }
            )
        elif risk_rows == 0:
            dossier_diagnostics.append("structured_risk_table_absent")
            diagnostics.append(
                {
                    "code": "structured_risk_table_absent",
                    "dossier_id": dossier.dossier_id,
                    "posture": "prose_only_not_synthesized",
                    "message": "No structured risk rows were extracted; prose risk language remains only in the local DOCX.",
                }
            )
        if dossier.metadata_identity_posture != "docx_metadata_cross_checked":
            dossier_diagnostics.append(dossier.metadata_identity_posture)
        route = PACKAGE_ROUTES[table_id].get(dossier.dossier_id)
        route_fields = (
            explicit_route_fields(
                dossier.dossier_id,
                master_status=dossier.master_status,
                master_confidence=dossier.master_confidence,
            )
            if route
            else {}
        )
        dossier_rows.append(
            {
                **({"admission_status": dossier.admission_status} if table_id != "table-i" else {}),
                "dossier_id": dossier.dossier_id,
                "docx_section": dossier.docx_section,
                "metadata_headers": dossier.metadata_headers,
                "metadata_identity_posture": dossier.metadata_identity_posture,
                "review_posture": route_fields.get("review_posture", str(route["review_posture"])) if route else "not_applicable_quarantined",
                "route_kind": str(route["route_kind"]) if route else "blocked_identity_mismatch",
                "structured_risk_row_count": risk_rows,
                "coverage": summary,
                "diagnostics": dossier_diagnostics,
            }
        )
    by_section = {
        section: coverage_summary([dossier for dossier in dossiers if dossier.docx_section == section])
        for section in PACKAGES[table_id].get("docx_sections", [])
    }
    return {
        "schema_version": "tos_philosophy_docx_extraction_coverage_v1",
        "path": repo_ref(extraction_coverage),
        "owner_repo": "Tree-of-Sophia",
        "owner_surface": "ToS/philosophy/graph-workbench/PLANTING_INTERFACE.md",
        **({"table_id": table_id} if table_id != "table-i" else {}),
        "intake_manifest_ref": repo_ref(intake_manifest),
        "route_map_ref": repo_ref(PREPARED_DOSSIER_ROUTES),
        "coverage_posture": "bounded_structured_planting_not_full_dossier_transfer",
        "structured_primary_families": sorted(STRUCTURED_FAMILIES),
        "identity_metadata_rule": (
            "Поле|Значение rows are examined only to cross-check Table I and ROW_TO_EXPAND; other metadata values "
            "are not represented as full dossier transfer. The A44 Параметр|Значение alias is also inspected for "
            "those two identity fields, but the table remains counted as deferred context."
            if table_id == "table-i"
            else "Поле|Значение, Поле|Идентификация, and Параметр|Значение rows are examined only to cross-check "
            "the package table and ROW_TO_EXPAND identity; other metadata values are not represented as full dossier transfer."
        ),
        "deferred_context_rule": (
            "Context table rows remain in the operator-local DOCX and are counted here; they are not silently "
            "promoted into nodes, relations, source anchors, terms, transmissions, source witnesses, or canon."
        ),
        "summary": coverage_summary(dossiers),
        "sections": by_section,
        "dossiers": dossier_rows,
        "diagnostics": diagnostics,
        "claim_limit": (
            "Coverage accounts for non-empty DOCX table-body rows under the current parser. It does not cover prose "
            "paragraph semantics, verify citations, accept risk judgments, or establish complete dossier transfer."
        ),
    }


def write_intake_and_coverage_surfaces(dossiers: list[Dossier]) -> None:
    table_id = dossiers[0].table_id
    write_json(package_path(table_id, "intake_manifest_ref"), build_intake_manifest(dossiers))
    write_json(package_path(table_id, "extraction_coverage_ref"), build_extraction_coverage(dossiers))


def update_atlas(dossiers: list[Dossier]) -> None:
    admitted = [dossier for dossier in dossiers if dossier.admission_status == "admitted"]
    by_table = {table_id: [dossier for dossier in dossiers if dossier.table_id == table_id] for table_id in SUPPORTED_TABLES}
    for table_id, package_dossiers in by_table.items():
        admitted_by_id = {
            dossier.dossier_id: dossier
            for dossier in package_dossiers
            if dossier.admission_status == "admitted"
        }
        blocked_by_id = {
            dossier.dossier_id: dossier
            for dossier in package_dossiers
            if dossier.admission_status != "admitted"
        }
        rows = load_jsonl(master_rows_path(table_id))
        for row in rows:
            row_id = str(row.get("row_id") or "")
            normalized = row.setdefault("normalized", {})
            if row_id in admitted_by_id:
                row["dossier_available"] = True
                row["dossier_id"] = row_id
                normalized["prepared_branch_path"] = str(admitted_by_id[row_id].branch_path)
                if table_id != "table-i":
                    normalized["dossier_intake_status"] = "admitted"
                else:
                    normalized.pop("dossier_intake_status", None)
            else:
                row["dossier_available"] = False
                row["dossier_id"] = None
                normalized.pop("prepared_branch_path", None)
                if row_id in blocked_by_id:
                    normalized["dossier_intake_status"] = blocked_by_id[row_id].admission_status
                elif row_id in PACKAGES[table_id].get("missing_master_dossier_ids", []):
                    normalized["dossier_intake_status"] = "input_not_supplied"
                else:
                    normalized.pop("dossier_intake_status", None)
        write_jsonl(master_rows_path(table_id), rows)

        table_manifest = load_json(master_manifest_path(table_id))
        table_manifest["available_dossiers"] = sorted(admitted_by_id)
        if table_id != "table-i":
            table_manifest["available_dossier_count"] = len(admitted_by_id)
            table_manifest["observed_input_count"] = len(package_dossiers)
            table_manifest["quarantined_dossiers"] = sorted(blocked_by_id)
            table_manifest["missing_master_dossier_ids"] = list(PACKAGES[table_id].get("missing_master_dossier_ids", []))
        else:
            for key in (
                "available_dossier_count",
                "observed_input_count",
                "quarantined_dossiers",
                "missing_master_dossier_ids",
            ):
                table_manifest.pop(key, None)
        write_json(master_manifest_path(table_id), table_manifest)

    atlas_manifest = load_json(ATLAS_MANIFEST)
    atlas_manifest["dossiers"]["available_count"] = len(admitted)
    atlas_manifest["dossiers"]["route_map"] = repo_ref(PREPARED_DOSSIER_ROUTES)
    atlas_manifest["dossiers"]["intake_manifests"] = [
        repo_ref(package_path(table_id, "intake_manifest_ref")) for table_id in SUPPORTED_TABLES
    ]
    atlas_manifest["dossiers"]["extraction_coverages"] = [
        repo_ref(package_path(table_id, "extraction_coverage_ref")) for table_id in SUPPORTED_TABLES
    ]
    atlas_manifest["dossiers"]["intake_manifest"] = repo_ref(INTAKE_MANIFEST)
    atlas_manifest["dossiers"]["extraction_coverage"] = repo_ref(EXTRACTION_COVERAGE)
    atlas_manifest["role"] = (
        "prepared atlas for ToS philosophy growth from master tables and admitted prepared dossiers"
    )
    write_json(ATLAS_MANIFEST, atlas_manifest)

    dossier_branch = load_json(DOSSIER_BRANCH_MANIFEST)
    dossier_branch["dossier_count"] = len(admitted)
    dossier_branch["source_anchor_backlog"] = repo_ref(SOURCE_ANCHOR_BACKLOG)
    dossier_branch["term_index"] = repo_ref(TERM_INDEX)
    dossier_branch["transmission_backlog"] = repo_ref(TRANSMISSION_BACKLOG)
    dossier_branch["prepared_dossier_routes"] = repo_ref(PREPARED_DOSSIER_ROUTES)
    dossier_branch["intake_manifests"] = atlas_manifest["dossiers"]["intake_manifests"]
    dossier_branch["extraction_coverages"] = atlas_manifest["dossiers"]["extraction_coverages"]
    dossier_branch["intake_manifest"] = repo_ref(INTAKE_MANIFEST)
    dossier_branch["extraction_coverage"] = repo_ref(EXTRACTION_COVERAGE)
    dossier_branch["role"] = (
        "index of admitted prepared Deep Research dossiers and their graph-shape tables"
    )
    write_json(DOSSIER_BRANCH_MANIFEST, dossier_branch)


def write_dossier_indexes(dossiers: list[Dossier]) -> None:
    dossiers = [dossier for dossier in dossiers if dossier.admission_status == "admitted"]
    node_counter: Counter[str] = Counter()
    relation_counter: Counter[str] = Counter()
    index_rows: list[dict[str, Any]] = []
    all_sources: list[dict[str, Any]] = []
    all_terms: list[dict[str, Any]] = []
    all_transmissions: list[dict[str, Any]] = []
    for dossier in dossiers:
        route = route_metadata(dossier.dossier_id)
        route_fields = explicit_route_fields(
            dossier.dossier_id,
            master_status=dossier.master_status,
            master_confidence=dossier.master_confidence,
        )
        node_type_counts = Counter(str(row.get("node_kind") or "unspecified") for row in dossier.node_rows)
        relation_counts = Counter(str(row.get("relation_kind") or "related_to") for row in dossier.relation_rows)
        node_counter.update(node_type_counts)
        relation_counter.update(relation_counts)
        index_rows.append(
            {
                "atlas_status": "prepared_dossier_indexed",
                "branch_path": dossier.branch_path,
                "dossier_id": dossier.dossier_id,
                "docx_section": dossier.docx_section,
                "extraction_coverage_ref": repo_ref(package_path(dossier.table_id, "extraction_coverage_ref")),
                "intake_manifest_ref": repo_ref(package_path(dossier.table_id, "intake_manifest_ref")),
                "master_table": dossier.master_table,
                "table_id": dossier.table_id,
                "master_confidence": dossier.master_confidence,
                "master_status": dossier.master_status,
                "metadata_identity_posture": dossier.metadata_identity_posture,
                "node_row_count": len(dossier.node_rows),
                "node_type_counts": dict(sorted(node_type_counts.items())),
                "paragraph_count": dossier.paragraph_count,
                "relation_counts": dict(sorted(relation_counts.items())),
                "relation_row_count": len(dossier.relation_rows),
                "source_anchor_count": len(dossier.source_rows),
                "source_document": dossier.source_document,
                "table_count": dossier.table_count,
                "table_row": dossier.table_row,
                "term_count": len(dossier.term_rows),
                "title": dossier.title,
                "transmission_count": len(dossier.transmission_rows),
                "review_posture": str(route_fields.get("review_posture") or route["review_posture"]),
                **(
                    {"review_reason": str(route_fields.get("review_reason") or route.get("review_reason"))}
                    if route_fields.get("review_reason") or route.get("review_reason")
                    else {}
                ),
                "route_kind": str(route["route_kind"]),
                **(
                    {"route_constraints": [str(value) for value in route["route_constraints"]]}
                    if isinstance(route.get("route_constraints"), list)
                    else {}
                ),
            }
        )
        all_sources.extend(dossier.source_rows)
        all_terms.extend(dossier.term_rows)
        all_transmissions.extend(dossier.transmission_rows)

    write_jsonl(DOSSIER_INDEX, sorted(index_rows, key=lambda row: (str(row["table_id"]), str(row["dossier_id"]))))
    write_json(
        DOSSIER_SUMMARY,
        {
            "schema_version": "tos_philosophy_atlas_dossier_graph_shape_v1",
            "path": repo_ref(DOSSIER_SUMMARY),
            "source": "supported prepared Deep Research dossier DOCX packages",
            "source_posture": "bounded structured extraction from operator-local non-authoritative research packets",
            "intake_manifest_ref": repo_ref(INTAKE_MANIFEST),
            "intake_manifest_refs": [repo_ref(package_path(table_id, "intake_manifest_ref")) for table_id in SUPPORTED_TABLES],
            "extraction_coverage_ref": repo_ref(EXTRACTION_COVERAGE),
            "extraction_coverage_refs": [repo_ref(package_path(table_id, "extraction_coverage_ref")) for table_id in SUPPORTED_TABLES],
            "dossier_count": len(dossiers),
            "node_row_count": sum(len(dossier.node_rows) for dossier in dossiers),
            "relation_row_count": sum(len(dossier.relation_rows) for dossier in dossiers),
            "node_type_counts": dict(sorted(node_counter.items())),
            "relation_counts": dict(sorted(relation_counter.items())),
            "source_anchor_count": len(all_sources),
            "term_count": len(all_terms),
            "transmission_count": len(all_transmissions),
        },
    )
    write_jsonl(SOURCE_ANCHOR_BACKLOG, all_sources)
    write_jsonl(TERM_INDEX, all_terms)
    write_jsonl(TRANSMISSION_BACKLOG, all_transmissions)


def dossier_local_node_alias(dossier_id: str, original_node_id: str) -> str | None:
    value = original_node_id.strip()
    prefixes = (dossier_id, dossier_id.replace("-", ""))
    for prefix in prefixes:
        if value.startswith(prefix) and len(value) > len(prefix) and value[len(prefix)] in ".-:":
            alias = value[len(prefix) + 1 :].strip()
            return alias or None
    return None


def resolve_dossier_endpoint(
    label: str,
    exact_lookup: dict[str, str],
    casefold_lookup: dict[str, str],
    local_alias_lookup: dict[str, str],
) -> str | None:
    value = label.strip()
    exact_candidate = exact_lookup.get(value)
    if exact_candidate:
        return exact_candidate
    casefold_candidate = casefold_lookup.get(value.casefold())
    if casefold_candidate:
        return casefold_candidate
    for alias in sorted(local_alias_lookup, key=lambda item: (-len(item), item)):
        if not value.startswith(alias) or len(value) == len(alias):
            continue
        boundary = value[len(alias)]
        if boundary.isspace() or boundary in ":;,.—–-":
            return local_alias_lookup[alias]
    return None


def resolve_relation_endpoints(dossiers: list[Dossier]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    all_nodes: list[dict[str, Any]] = []
    all_relations: list[dict[str, Any]] = []
    nodes_by_dossier: dict[str, dict[str, str]] = {}
    casefold_nodes_by_dossier: dict[str, dict[str, str]] = {}
    local_aliases_by_dossier: dict[str, dict[str, str]] = {}
    for dossier in dossiers:
        lookup: dict[str, str] = {}
        casefold_claims: dict[str, set[str]] = {}
        local_alias_claims: dict[str, set[str]] = {}
        for row in dossier.node_rows:
            all_nodes.append(row)
            for key in ("original_node_id", "label"):
                value = str(row.get(key) or "").strip()
                if value:
                    lookup[value] = str(row["candidate_id"])
                    casefold_claims.setdefault(value.casefold(), set()).add(str(row["candidate_id"]))
                if key == "original_node_id":
                    local_alias = dossier_local_node_alias(dossier.dossier_id, value)
                    if local_alias:
                        local_alias_claims.setdefault(local_alias, set()).add(str(row["candidate_id"]))
        for local_alias, candidate_ids in local_alias_claims.items():
            if len(candidate_ids) == 1 and local_alias not in lookup:
                lookup[local_alias] = next(iter(candidate_ids))
        nodes_by_dossier[dossier.dossier_id] = lookup
        casefold_nodes_by_dossier[dossier.dossier_id] = {
            value: next(iter(candidate_ids))
            for value, candidate_ids in casefold_claims.items()
            if len(candidate_ids) == 1
        }
        local_aliases_by_dossier[dossier.dossier_id] = {
            local_alias: next(iter(candidate_ids))
            for local_alias, candidate_ids in local_alias_claims.items()
            if len(candidate_ids) == 1
        }
    for dossier in dossiers:
        lookup = nodes_by_dossier[dossier.dossier_id]
        casefold_lookup = casefold_nodes_by_dossier[dossier.dossier_id]
        local_alias_lookup = local_aliases_by_dossier[dossier.dossier_id]
        for row in dossier.relation_rows:
            source = str(row.get("source_endpoint_label") or "")
            target = str(row.get("target_endpoint_label") or "")
            row["source_candidate_id"] = resolve_dossier_endpoint(source, lookup, casefold_lookup, local_alias_lookup)
            row["target_candidate_id"] = resolve_dossier_endpoint(target, lookup, casefold_lookup, local_alias_lookup)
            row["endpoint_resolution"] = (
                "matched_nodes" if row["source_candidate_id"] and row["target_candidate_id"] else "label_endpoint"
            )
            all_relations.append(row)
    return all_nodes, all_relations


def _text_status_for_ru(label: str) -> tuple[str, str]:
    translated, status = russian_label(label)
    if re.search(r"[А-Яа-яЁё]", label):
        return translated, status
    if translated != label or status == "reviewed":
        return translated, status
    return translated, "pending"


def _text_status_for_en(label: str) -> tuple[str, str]:
    translated, status = english_label(label)
    if status == "draft" and translated:
        translated = translated[0].upper() + translated[1:]
    return translated, status if translated != label or status != "source" else "source"


def extracted_row_table_id(row: dict[str, Any]) -> str:
    value = row.get("table_id")
    if isinstance(value, str) and value:
        return value
    return "table-ii" if str(row.get("dossier_id") or "").startswith("T2-") else "table-i"


def text_bearing_language_packet(row: dict[str, Any]) -> dict[str, Any]:
    label = str(row.get("label") or "").strip()
    ru, ru_status = _text_status_for_ru(label)
    en, en_status = _text_status_for_en(label)
    table_id = extracted_row_table_id(row)
    source_ref = repo_ref(language_packets_path(table_id))
    source_node_ref = str(row.get("source_ref") or repo_ref(proposed_nodes_path(table_id)))
    source_refs = sorted({source_ref, source_node_ref, LANGUAGE_PACKET_CONTRACT_REF, LANGUAGE_REGISTRY_REF})
    review_metadata = {
        field: row[field]
        for field in ("review_posture", "review_reason", "master_status", "master_confidence")
        if row.get(field) is not None
    }
    return {
        "schema_version": "tos_philosophy_text_bearing_language_packet_v1",
        "packet_id": f"language-packet:{row['candidate_id']}",
        "node_ref": {
            "id": row["candidate_id"],
            "id_kind": "candidate_id",
            "source_ref": source_node_ref,
        },
        "node_kind": row.get("node_kind"),
        "branch_path": row.get("branch_path"),
        "authority_posture": row.get("authority_posture"),
        "canon_status": row.get("canon_status"),
        "atlas_row_id": row.get("atlas_row_id"),
        "dossier_id": row.get("dossier_id"),
        **({"table_id": table_id} if table_id != "table-i" else {}),
        **review_metadata,
        "source_document": row.get("source_document"),
        "source_row_index": row.get("source_row_index"),
        "source_table_index": row.get("source_table_index"),
        "source_label": label,
        "source_ref": source_ref,
        "source_refs": source_refs,
        "language_registry_ref": LANGUAGE_REGISTRY_REF,
        "text_bearing_nodes_contract_ref": LANGUAGE_PACKET_CONTRACT_REF,
        "title_block": {
            "original": {
                "value": None,
                "language": "und",
                "script": "Zzzz",
                "transliteration": None,
                "attestation_status": "unknown",
                "source_ref": source_node_ref,
                "review_status": "pending_original_witness",
            },
            "ru": {
                "value": ru,
                "translation_status": ru_status,
                "source_ref": source_node_ref,
            },
            "en": {
                "value": en,
                "translation_status": en_status,
                "source_ref": source_node_ref,
            },
        },
        "witness_block": {
            "source_witnesses": [],
            "witness_status": "pending_source_witness_anchor",
            "required_next_fields": [
                "witness_id",
                "witness_kind",
                "language",
                "script",
                "repository_or_corpus",
                "source_ref",
            ],
        },
        "relation_pressure": [
            {
                "predicate": "has_original_language",
                "target_status": "unresolved",
                "review_route": LANGUAGE_PACKET_CONTRACT_REF,
            },
            {
                "predicate": "uses_script",
                "target_status": "unresolved",
                "review_route": LANGUAGE_PACKET_CONTRACT_REF,
            },
            {
                "predicate": "has_witness",
                "target_status": "unresolved",
                "review_route": "ToS/philosophy/atlas/dossiers/source-anchor-backlog.jsonl",
            },
        ],
    }


def build_language_packets(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        text_bearing_language_packet(row)
        for row in nodes
        if str(row.get("node_kind") or "") in TEXT_BEARING_NODE_KINDS
    ]


def write_language_packet_surfaces(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    packets = build_language_packets(nodes)
    packet_refs: list[str] = []
    for table_id in SUPPORTED_TABLES:
        table_packets = [packet for packet in packets if extracted_row_table_id(packet) == table_id]
        path = language_packets_path(table_id)
        write_jsonl(path, table_packets)
        packet_refs.append(repo_ref(path))
    write_json(
        LANGUAGE_PACKETS_MANIFEST,
        {
            "branch_id": "philosophy.graph-workbench.language-packets",
            "path": "ToS/philosophy/graph-workbench/language-packets",
            "role": "text-bearing language packets emitted by prepared dossier planting before canon promotion",
            "contract_ref": LANGUAGE_PACKET_CONTRACT_REF,
            "language_registry_ref": LANGUAGE_REGISTRY_REF,
            "packet_count": len(packets),
            "source_ref": repo_ref(LANGUAGE_PACKETS),
            "source_refs": packet_refs,
        },
    )
    table_rows = "".join(
        f"| `{table_id}-text-bearing-nodes.jsonl` | generated packets for admitted {table_id} text-corpus candidates |\n"
        for table_id in SUPPORTED_TABLES
    )
    LANGUAGE_PACKETS_README.write_text(
        "# Language Packets\n\n"
        "`language-packets/` contains pre-canon language packets for text-bearing philosophy nodes.\n\n"
        "The package files are generated from supported prepared-dossier proposed nodes. "
        "It records original-language uncertainty, Russian and English display slots, witness posture, "
        "and language/script relation pressure without promoting any source claim to canon.\n\n"
        "| Surface | Role |\n"
        "| --- | --- |\n"
        f"{table_rows}"
        f"| `{LANGUAGE_PACKET_CONTRACT_REF}` | packet contract |\n"
        f"| `{LANGUAGE_REGISTRY_REF}` | language and script registry |\n",
        encoding="utf-8",
    )
    return packets


def write_graph_workbench(dossiers: list[Dossier]) -> None:
    dossiers = [dossier for dossier in dossiers if dossier.admission_status == "admitted"]
    nodes, relations = resolve_relation_endpoints(dossiers)
    language_packets = write_language_packet_surfaces(nodes)
    for table_id in SUPPORTED_TABLES:
        table_dossiers = [dossier for dossier in dossiers if dossier.table_id == table_id]
        table_nodes = [row for row in nodes if extracted_row_table_id(row) == table_id]
        table_relations = [row for row in relations if extracted_row_table_id(row) == table_id]
        table_packets = [packet for packet in language_packets if extracted_row_table_id(packet) == table_id]
        nodes_path = proposed_nodes_path(table_id)
        relations_path = proposed_relations_path(table_id)
        fragments_path = branch_fragments_path(table_id)
        ledger_path = promotion_ledger_path(table_id)
        write_jsonl(nodes_path, table_nodes)
        write_jsonl(relations_path, table_relations)
        branches: list[dict[str, Any]] = []
        for dossier in sorted(table_dossiers, key=dossier_sort_key):
            route = route_metadata(dossier.dossier_id)
            route_fields = explicit_route_fields(
                dossier.dossier_id,
                master_status=dossier.master_status,
                master_confidence=dossier.master_confidence,
            )
            branches.append(
                {
                    "branch_path": dossier.branch_path,
                    "dossier_id": dossier.dossier_id,
                    "docx_section": dossier.docx_section,
                    "title": dossier.title,
                    "node_row_count": len(dossier.node_rows),
                    "relation_row_count": len(dossier.relation_rows),
                    "source_anchor_count": len(dossier.source_rows),
                    "term_count": len(dossier.term_rows),
                    "transmission_count": len(dossier.transmission_rows),
                    "review_posture": str(route_fields.get("review_posture") or route["review_posture"]),
                    "route_kind": str(route["route_kind"]),
                }
            )
        write_json(
            fragments_path,
            {
                "schema_version": "tos_philosophy_branch_fragments_v1",
                "path": repo_ref(fragments_path),
                "source_ref": repo_ref(DOSSIER_INDEX),
                **({"table_id": table_id} if table_id != "table-i" else {}),
                "canon_status": "pre-canon",
                "branch_count": len(table_dossiers),
                "language_packet_count": len(table_packets),
                "language_packets_ref": repo_ref(language_packets_path(table_id)),
                "branches": branches,
            },
        )
        ledger_path.parent.mkdir(parents=True, exist_ok=True)
        table_title = {"table-i": "Table I", "table-ii": "Table II", "table-iii": "Table III"}[table_id]
        intro = f"This ledger records the bounded planting of the admitted {table_title} dossier package."
        ledger_path.write_text(
            f"# {table_title} Prepared Dossiers\n\n"
            f"{intro}\n\n"
            "| Surface | Count | Status |\n"
            "| --- | ---: | --- |\n"
            f"| prepared dossiers | {len(table_dossiers)} | atlas indexed |\n"
            f"| proposed nodes | {len(table_nodes)} | pre-canon graph workbench |\n"
            f"| proposed relations | {len(table_relations)} | pre-canon graph workbench |\n"
            f"| text-bearing language packets | {len(table_packets)} | pre-canon multilingual review |\n"
            f"| branch fragments | {len(table_dossiers)} | era/region/tradition or explicit frontier branch bodies |\n\n"
            "Promotion remains a later authored review step through ToS canon route cards.\n",
            encoding="utf-8",
        )


def branch_title(dossier: Dossier) -> str:
    if dossier.table_id == "table-i":
        return dossier.title.replace("ToS Deep Research:", "").strip()
    reviewed, _status = russian_label(dossier.dossier_id)
    prefix = f"{dossier.dossier_id} — "
    return reviewed.removeprefix(prefix)


def top_counts(rows: list[dict[str, Any]], key: str, limit: int = 8) -> list[str]:
    counter = Counter(str(row.get(key) or "unspecified") for row in rows)
    return [f"{name}: {count}" for name, count in counter.most_common(limit)]


def render_branch_readme(dossier: Dossier) -> str:
    title = branch_title(dossier)
    node_pressure = ", ".join(top_counts(dossier.node_rows, "node_kind")) or "none"
    relation_pressure = ", ".join(top_counts(dossier.relation_rows, "relation_kind")) or "none"
    path_ref = str(dossier.branch_path)
    route = route_metadata(dossier.dossier_id)
    route_fields = explicit_route_fields(
        dossier.dossier_id,
        master_status=dossier.master_status,
        master_confidence=dossier.master_confidence,
    )
    planting_paths = sorted(
        (REPO_ROOT / path_ref / "sources/plantings").glob(
            "*/source-planting.json"
        )
    )
    planting_row = (
        "| `sources/plantings/` | exact source-witness routes already planted from backlog anchors |\n"
        if planting_paths
        else ""
    )
    review_posture = str(route_fields.get("review_posture") or route["review_posture"])
    review_reason = str(route_fields.get("review_reason") or route.get("review_reason") or "")
    route_constraints = route.get("route_constraints")
    review_block = (
        "## Manual Review Gate\n\n"
        f"- Posture: `{review_posture}`\n"
        f"- Master-table status/confidence: `{dossier.master_status}/{dossier.master_confidence}`\n"
        f"- Reason: {review_reason}\n\n"
        if review_posture == "manual_review_required"
        else ""
    )
    constraint_block = (
        "## Route Constraints\n\n"
        + "".join(f"- `{value}`\n" for value in route_constraints)
        + "\n"
        if isinstance(route_constraints, list) and route_constraints
        else ""
    )
    return (
        f"# {title}\n\n"
        f"Atlas row: `{dossier.dossier_id}`. Prepared dossier: `{dossier.source_document}`.\n\n"
        "This branch is the ToS philosophy home for the prepared dossier's first tree-shaped growth. "
        "It keeps the dossier material in the philosophy tree while leaving canon promotion to a later authored review.\n\n"
        "## Branch Pressure\n\n"
        f"- Candidate node rows: {len(dossier.node_rows)}\n"
        f"- Candidate relation rows: {len(dossier.relation_rows)}\n"
        f"- Source-anchor backlog rows: {len(dossier.source_rows)}\n"
        f"- Term rows: {len(dossier.term_rows)}\n"
        f"- Transmission rows: {len(dossier.transmission_rows)}\n"
        f"- Node pressure: {node_pressure}\n"
        f"- Relation pressure: {relation_pressure}\n\n"
        f"{review_block}"
        f"{constraint_block}"
        "## Local Surfaces\n\n"
        "| Surface | Role |\n"
        "| --- | --- |\n"
        "| `sources/source-anchor-backlog.jsonl` | future real source witness and edition anchors for this branch |\n"
        f"{planting_row}"
        "| `graph-workbench/pre-canon-summary.json` | local summary of proposed graph rows before canon review |\n\n"
        "Global proposed node and relation rows for this branch are aggregated in "
        f"`{repo_ref(proposed_nodes_path(dossier.table_id))}` and "
        f"`{repo_ref(proposed_relations_path(dossier.table_id))}`.\n"
    )


def remove_obsolete_generated_branch() -> None:
    if not OBSOLETE_GENERATED_BRANCH.exists():
        return
    files = [path for path in OBSOLETE_GENERATED_BRANCH.rglob("*") if path.is_file()]
    if any(path.name != "branch.manifest.json" for path in files):
        raise SystemExit(f"{repo_ref(OBSOLETE_GENERATED_BRANCH)} contains non-generated files; refusing to remove it")
    shutil.rmtree(OBSOLETE_GENERATED_BRANCH)


def write_branch_surfaces(dossiers: list[Dossier]) -> None:
    dossiers = [dossier for dossier in dossiers if dossier.admission_status == "admitted"]
    remove_obsolete_generated_branch()
    parent_children: dict[str, set[str]] = defaultdict(set)
    parent_roles: dict[str, str] = {}
    for package in PACKAGES.values():
        ancestor_roles = package.get("ancestor_roles")
        if isinstance(ancestor_roles, dict):
            for ancestor_path, ancestor_role in ancestor_roles.items():
                parent_roles[str(ancestor_path)] = str(ancestor_role)
    for dossier_id, (path_ref, _role) in ALL_BRANCHES.items():
        ancestor_roles = route_metadata(dossier_id).get("ancestor_roles")
        if isinstance(ancestor_roles, dict):
            for ancestor_path, ancestor_role in ancestor_roles.items():
                if ancestor_path in parent_roles and parent_roles[ancestor_path] != ancestor_role:
                    raise ValueError(f"conflicting ancestor role for {ancestor_path}")
                parent_roles[str(ancestor_path)] = str(ancestor_role)
        parts = Path(path_ref).parts
        for index in range(3, len(parts)):
            parent = Path(*parts[:index]).as_posix()
            child = Path(*parts[index : index + 1]).as_posix()
            parent_children[parent].add(child)

    existing_manifest_paths = {
        repo_ref(path): load_json(path)
        for path in (REPO_ROOT / "ToS/philosophy").glob("**/branch.manifest.json")
        if not path.is_relative_to(OBSOLETE_GENERATED_BRANCH)
    }

    for path_ref, children in sorted(parent_children.items()):
        path = REPO_ROOT / path_ref
        path.mkdir(parents=True, exist_ok=True)
        existing = existing_manifest_paths.get(f"{path_ref}/branch.manifest.json", {})
        existing_children = {
            child for child in existing.get("children", []) if isinstance(child, str) and child != "regions/ancient-near-east"
        }
        payload = {
            "branch_id": branch_id_for(path_ref),
            "children": sorted(existing_children | {child for child in children}),
            "path": path_ref,
            "role": existing.get("role") or parent_roles.get(path_ref) or f"{Path(path_ref).name} philosophy branch",
        }
        write_json(path / "branch.manifest.json", payload)

    for dossier in dossiers:
        path_ref = str(dossier.branch_path)
        role = str(dossier.branch_role)
        path = REPO_ROOT / path_ref
        path.mkdir(parents=True, exist_ok=True)
        planting_refs = sorted(
            repo_ref(planting_path)
            for planting_path in (path / "sources/plantings").glob(
                "*/source-planting.json"
            )
        )
        write_json(
            path / "branch.manifest.json",
            {
                "branch_id": branch_id_for(path_ref),
                "path": path_ref,
                "role": role,
                "atlas_rows": [dossier.dossier_id],
                "prepared_dossiers": [dossier.dossier_id],
                "evidence_status": "prepared_dossier_branch",
                **explicit_route_fields(
                    dossier.dossier_id,
                    master_status=dossier.master_status,
                    master_confidence=dossier.master_confidence,
                ),
                "expected_local_children": ["sources", "graph-workbench"],
                "source_anchor_backlog": f"{path_ref}/sources/source-anchor-backlog.jsonl",
                **(
                    {
                        "source_planting_count": len(planting_refs),
                        "source_planting_refs": planting_refs,
                    }
                    if planting_refs
                    else {}
                ),
                "local_graph_summary": f"{path_ref}/graph-workbench/pre-canon-summary.json",
            },
        )
        (path / "README.md").write_text(render_branch_readme(dossier), encoding="utf-8")

        sources_path = path / "sources"
        graph_path = path / "graph-workbench"
        sources_path.mkdir(exist_ok=True)
        graph_path.mkdir(exist_ok=True)
        write_json(
            sources_path / "branch.manifest.json",
            {
                "branch_id": branch_id_for(f"{path_ref}/sources"),
                "path": f"{path_ref}/sources",
                "role": (
                    f"source-anchor backlog and planted source routes for {dossier.dossier_id}"
                    if planting_refs
                    else f"source-anchor backlog for {dossier.dossier_id}"
                ),
                "anchor_count": len(dossier.source_rows),
                **explicit_route_fields(
                    dossier.dossier_id,
                    master_status=dossier.master_status,
                    master_confidence=dossier.master_confidence,
                ),
                **(
                    {
                        "planting_count": len(planting_refs),
                        "planting_refs": planting_refs,
                    }
                    if planting_refs
                    else {}
                ),
            },
        )
        write_jsonl(sources_path / "source-anchor-backlog.jsonl", dossier.source_rows)
        write_json(
            graph_path / "branch.manifest.json",
            {
                "branch_id": branch_id_for(f"{path_ref}/graph-workbench"),
                "path": f"{path_ref}/graph-workbench",
                "role": f"local pre-canon graph summary for {dossier.dossier_id}",
                "proposed_node_count": len(dossier.node_rows),
                "proposed_relation_count": len(dossier.relation_rows),
                **explicit_route_fields(
                    dossier.dossier_id,
                    master_status=dossier.master_status,
                    master_confidence=dossier.master_confidence,
                ),
            },
        )
        write_json(
            graph_path / "pre-canon-summary.json",
            {
                "schema_version": "tos_philosophy_local_graph_summary_v1",
                "path": f"{path_ref}/graph-workbench/pre-canon-summary.json",
                "atlas_row_id": dossier.dossier_id,
                "dossier_id": dossier.dossier_id,
                "branch_path": path_ref,
                "canon_status": "pre-canon",
                **explicit_route_fields(
                    dossier.dossier_id,
                    master_status=dossier.master_status,
                    master_confidence=dossier.master_confidence,
                ),
                "proposed_nodes_ref": repo_ref(proposed_nodes_path(dossier.table_id)),
                "proposed_relations_ref": repo_ref(proposed_relations_path(dossier.table_id)),
                "node_row_count": len(dossier.node_rows),
                "relation_row_count": len(dossier.relation_rows),
                "node_type_counts": dict(sorted(Counter(str(row.get("node_kind")) for row in dossier.node_rows).items())),
                "relation_counts": dict(sorted(Counter(str(row.get("relation_kind")) for row in dossier.relation_rows).items())),
            },
        )


def refresh_philosophy_manifest() -> None:
    manifest = load_json(PHILOSOPHY_MANIFEST)
    manifest["branch_manifests"] = sorted(
        repo_ref(path)
        for path in (REPO_ROOT / "ToS/philosophy").glob("**/branch.manifest.json")
    )
    atlas_routes = set(manifest.get("atlas_routes", []))
    atlas_routes.update(
        {
            repo_ref(SOURCE_ANCHOR_BACKLOG),
            repo_ref(TERM_INDEX),
            repo_ref(TRANSMISSION_BACKLOG),
            repo_ref(PREPARED_DOSSIER_ROUTES),
            LANGUAGE_REGISTRY_REF,
            LANGUAGE_PACKET_CONTRACT_REF,
        }
    )
    manifest["atlas_routes"] = sorted(atlas_routes)
    research_packet_contracts = set(manifest.get("research_packet_contracts", []))
    research_packet_contracts.update(
        repo_ref(package_path(table_id, field))
        for table_id in SUPPORTED_TABLES
        for field in ("intake_manifest_ref", "extraction_coverage_ref")
    )
    manifest["research_packet_contracts"] = sorted(research_packet_contracts)
    write_json(PHILOSOPHY_MANIFEST, manifest)


def write_readmes() -> None:
    (REPO_ROOT / "ToS/philosophy/atlas/dossiers/README.md").write_text(
        "# Dossiers\n\n"
        "`dossiers/` indexes admitted prepared Deep Research documents for the philosophy atlas.\n\n"
        "The complete Table I, Table II, and Table III plantings record dossier identity, branch route, graph-row pressure, "
        "source-anchor backlog, terms, and transmission rows while keeping canon promotion separate.\n\n"
        "| Surface | Role |\n"
        "| --- | --- |\n"
        "| `index.jsonl` | one entry per admitted prepared dossier |\n"
        "| `graph-shape-summary.json` | aggregate node, relation, source-anchor, term, and transmission pressure |\n"
        "| `prepared-dossier-routes.json` | source-owned route map from prepared dossier ids to philosophy branch homes |\n"
        f"| `{repo_ref(INTAKE_MANIFEST)}` | tracked fixity and capture-posture manifest for the untracked local DOCX bytes |\n"
        f"| `{repo_ref(EXTRACTION_COVERAGE)}` | explicit structured extraction and deferred-context coverage |\n"
        f"| `{repo_ref(package_path('table-ii', 'intake_manifest_ref'))}` | Table II fixity and admission record |\n"
        f"| `{repo_ref(package_path('table-ii', 'extraction_coverage_ref'))}` | Table II structured and deferred row accounting |\n"
        f"| `{repo_ref(package_path('table-iii', 'intake_manifest_ref'))}` | Table III fixity and admission record |\n"
        f"| `{repo_ref(package_path('table-iii', 'extraction_coverage_ref'))}` | Table III structured and deferred row accounting |\n"
        "| `source-anchor-backlog.jsonl` | future source witness, edition, corpus, and risk-control anchors |\n"
        "| `term-index.jsonl` | prepared term rows extracted from dossier terminology tables |\n"
        "| `transmission-backlog.jsonl` | incoming and outgoing transmission rows extracted from dossier tables |\n\n"
        "Text-bearing language packets for graph review live in "
        "`ToS/philosophy/graph-workbench/language-packets/` and follow "
        "`ToS/philosophy/atlas/multilingual/text-bearing-nodes.contract.json`.\n\n"
        "Branch bodies live under `ToS/philosophy/eras/...` or an explicit `ToS/philosophy/frontiers/...` route, "
        "and pre-canon graph rows live under "
        "`ToS/philosophy/graph-workbench/`.\n",
        encoding="utf-8",
    )
    (REPO_ROOT / "ToS/philosophy/atlas/README.md").write_text(
        "# Philosophy Atlas\n\n"
        "`atlas/` is the prepared navigation body for the whole ToS philosophy tree.\n\n"
        "It holds the master-table row spine, admitted prepared-dossier index, and aggregate "
        "pressure maps that tell the philosophy tree what must grow next.\n\n"
        "## Shape\n\n"
        "```text\n"
        "atlas/\n"
        "  master-tables/\n"
        "    table-i/\n"
        "    table-ii/\n"
        "    table-iii/\n"
        "  dossiers/\n"
        "    index.jsonl\n"
        "    graph-shape-summary.json\n"
        "    source-anchor-backlog.jsonl\n"
        "    term-index.jsonl\n"
        "    transmission-backlog.jsonl\n"
        "  multilingual/\n"
        "    content-labels.json\n"
        "    language-registry.json\n"
        "    text-bearing-nodes.contract.json\n"
        "```\n\n"
        "The atlas is prepared navigation and growth pressure. Branch bodies live in "
        "`ToS/philosophy/eras/...` or the explicit non-era `ToS/philosophy/frontiers/...` route; pre-canon graph material lives in "
        "`ToS/philosophy/graph-workbench/`; authored canon relation packs live in the canon route.\n\n"
        "`multilingual/` is a source-owned display companion for the atlas and generated "
        "graph projections. It preserves the route rule that planted works must carry "
        "their attested original form when available, plus Russian and English display "
        "labels for review and downstream visualization.\n\n"
        "For works, corpora, inscriptions, source witnesses, translations, versions, and "
        "commentaries, `multilingual/text-bearing-nodes.contract.json` is the planting "
        "contract. It keeps original-language title posture, transliteration, Russian "
        "review text, English review/runtime text, witness posture, and graph relation "
        "pressure in one source-owned packet shape.\n",
        encoding="utf-8",
    )


def plant_supported_packages() -> int:
    dossiers: list[Dossier] = []
    dossiers_by_table: dict[str, list[Dossier]] = {}
    for table_id in SUPPORTED_TABLES:
        master_rows = {
            str(row.get("row_id")): row
            for row in load_jsonl(master_rows_path(table_id))
        }
        package_dossiers = [
            parse_dossier(path, master_rows[extract_dossier_id(path)], table_id)
            for path in discover_docx(table_id)
        ]
        package_dossiers.sort(key=dossier_sort_key)
        dossiers_by_table[table_id] = package_dossiers
        dossiers.extend(package_dossiers)

    for table_id in SUPPORTED_TABLES:
        write_intake_and_coverage_surfaces(dossiers_by_table[table_id])
    dossiers.sort(key=dossier_sort_key)
    update_atlas(dossiers)
    write_dossier_indexes(dossiers)
    write_graph_workbench(dossiers)
    write_branch_surfaces(dossiers)
    refresh_philosophy_manifest()
    write_readmes()
    print(
        "[ok] planted supported prepared dossiers: "
        f"{sum(dossier.admission_status == 'admitted' for dossier in dossiers)} admitted from {len(dossiers)} artifacts, "
        f"{sum(len(dossier.node_rows) for dossier in dossiers)} proposed nodes, "
        f"{sum(len(dossier.relation_rows) for dossier in dossiers)} proposed relations"
    )
    return 0


def main() -> int:
    from plant_prepared_dossiers import require_aggregate_readiness

    require_aggregate_readiness()
    return plant_supported_packages()


if __name__ == "__main__":
    raise SystemExit(main())
